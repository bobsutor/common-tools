# cspell:ignore yattag Aptos klass
# cspell:disable

from docx.shared import Pt

import docx_tools

# we have ommitted some if they can be confused with names in other countries

COUNTRY_DATA_FOR_NEWS_CATEGORIES = {
    # REGION: Africa
    "Algeria": {
        "cities": ["Algiers", "Constantine", "Oran"],
        "country-flag-emoji": "🇩🇿",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Angola": {
        "cities": ["Huambo", "Lobito", "Luanda"],
        "country-flag-emoji": "🇦🇴",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Benin": {
        "cities": ["Abomey-Calavi", "Cotonou", "Porto-Novo"],
        "country-flag-emoji": "🇧🇯",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Botswana": {
        "cities": ["Francistown", "Gaborone", "Molepolole"],
        "country-flag-emoji": "🇧🇼",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Burkina Faso": {
        "cities": ["Bobo-Dioulasso", "Koudougou", "Ouagadougou"],
        "country-flag-emoji": "🇧🇫",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Burundi": {
        "cities": ["Bujumbura", "Gitega", "Muyinga"],
        "country-flag-emoji": "🇧🇮",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Cabo Verde": {
        "cities": ["Mindelo", "Praia", "Santa Maria"],
        "country-flag-emoji": "🇨🇻",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Cameroon": {
        "cities": ["Douala", "Garoua", "Yaoundé"],
        "country-flag-emoji": "🇨🇲",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Central African Republic": {
        "cities": ["Bangui", "Berbérati", "Bimbo"],
        "country-flag-emoji": "🇨🇫",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Chad": {
        "cities": ["Moundou", "N'Djamena", "Sarh"],
        "country-flag-emoji": "🇹🇩",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Comoros": {
        "cities": ["Fomboni", "Moroni", "Mutsamudu"],
        "country-flag-emoji": "🇰🇲",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Democratic Republic of the Congo": {
        "cities": ["Kinshasa", "Lubumbashi", "Mbuji-Mayi"],
        "country-flag-emoji": "🇨🇩",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": ["Congo"],
        "regions": ["Africa"],
    },
    "Djibouti": {
        "cities": ["Ali Sabieh", "Dikhil", "Djibouti City"],
        "country-flag-emoji": "🇩🇯",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Egypt": {
        "cities": ["Alexandria", "Cairo", "Giza", "New Cairo"],
        "country-flag-emoji": "🇪🇬",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Egyptian"],
        "regions": ["Africa", "Middle East"],
    },
    "Equatorial Guinea": {
        "cities": ["Bata", "Ebebeyín", "Malabo"],
        "country-flag-emoji": "🇬🇶",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Eritrea": {
        "cities": ["Asmara", "Keren", "Massawa"],
        "country-flag-emoji": "🇪🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Eswatini": {
        "cities": ["Big Bend", "Manzini", "Mbabane"],
        "country-flag-emoji": "🇸🇿",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Ethiopia": {
        "cities": ["Addis Ababa", "Gondar", "Mekele"],
        "country-flag-emoji": "🇪🇹",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Gabon": {
        "cities": ["Franceville", "Libreville", "Port-Gentil"],
        "country-flag-emoji": "🇬🇦",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Gambia": {
        "cities": ["Bakau", "Brikama", "Serekunda"],
        "country-flag-emoji": "🇬🇲",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Ghana": {
        "cities": ["Accra", "Kumasi", "Tamale"],
        "country-flag-emoji": "🇬🇭",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Guinea": {
        "cities": ["Conakry", "Kankan", "Nzérékoré"],
        "country-flag-emoji": "🇬🇳",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Guinea-Bissau": {
        "cities": ["Bafatá", "Bissau", "Gabú"],
        "country-flag-emoji": "🇬🇼",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Ivory Coast": {
        "cities": ["Abidjan", "Bouaké", "Daloa"],
        "country-flag-emoji": "🇨🇮",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Kenya": {
        "cities": ["Kisumu", "Mombasa", "Nairobi"],
        "country-flag-emoji": "🇰🇪",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Lesotho": {
        "cities": ["Mafeteng", "Maseru", "Teyateyaneng"],
        "country-flag-emoji": "🇱🇸",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Liberia": {
        "cities": ["Gbarnga", "Kakata", "Monrovia"],
        "country-flag-emoji": "🇱🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Libya": {
        "cities": ["Benghazi", "Misrata", "Tripoli"],
        "country-flag-emoji": "🇱🇾",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Madagascar": {
        "cities": ["Antananarivo", "Antsirabe", "Toamasina"],
        "country-flag-emoji": "🇲🇬",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Malawi": {
        "cities": ["Blantyre", "Lilongwe", "Mzuzu"],
        "country-flag-emoji": "🇲🇼",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Mali": {
        "cities": ["Bamako", "Mopti", "Sikasso"],
        "country-flag-emoji": "🇲🇱",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Mauritania": {
        "cities": ["Kiffa", "Nouadhibou", "Nouakchott"],
        "country-flag-emoji": "🇲🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Mauritius": {
        "cities": ["Beau Bassin-Rose Hill", "Port Louis", "Vacoas-Phoenix"],
        "country-flag-emoji": "🇲🇺",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Morocco": {
        "cities": ["Casablanca", "Fez", "Tangier"],
        "country-flag-emoji": "🇲🇦",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Mozambique": {
        "cities": ["Beira", "Maputo", "Matola"],
        "country-flag-emoji": "🇲🇿",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Namibia": {
        "cities": ["Swakopmund", "Walvis Bay", "Windhoek"],
        "country-flag-emoji": "🇳🇦",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Niger": {
        "cities": ["Maradi", "Niamey", "Zinder"],
        "country-flag-emoji": "🇳🇪",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Nigeria": {
        "cities": ["Ibadan", "Kano", "Lagos"],
        "country-flag-emoji": "🇳🇬",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Republic of the Congo": {
        "cities": ["Brazzaville", "Dolisie", "Pointe-Noire"],
        "country-flag-emoji": "🇨🇬",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Rwanda": {
        "cities": ["Butare", "Gisenyi", "Kigali"],
        "country-flag-emoji": "🇷🇼",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Sao Tome and Principe": {
        "cities": ["Neves", "Santo Amaro", "São Tomé"],
        "country-flag-emoji": "🇸🇹",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Senegal": {
        "cities": ["Dakar", "Thiès", "Touba"],
        "country-flag-emoji": "🇸🇳",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Seychelles": {
        "cities": ["Anse Etoile", "Beau Vallon"],
        "country-flag-emoji": "🇸🇨",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Sierra Leone": {
        "cities": ["Freetown", "Kenema"],
        "country-flag-emoji": "🇸🇱",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Somalia": {
        "cities": ["Hargeisa", "Kismayo", "Mogadishu"],
        "country-flag-emoji": "🇸🇴",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "South Africa": {
        "cities": ["Cape Town", "Durban", "Johannesburg"],
        "country-flag-emoji": "🇿🇦",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "South Sudan": {
        "cities": ["Juba", "Wau", "Yei"],
        "country-flag-emoji": "🇸🇸",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Sudan": {
        "cities": ["Khartoum", "Khartoum North", "Omdurman"],
        "country-flag-emoji": "🇸🇩",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Tanzania": {
        "cities": ["Arusha", "Dar es Salaam", "Mwanza"],
        "country-flag-emoji": "🇹🇿",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Togo": {
        "cities": ["Kara", "Lomé", "Sokodé"],
        "country-flag-emoji": "🇹🇬",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Tunisia": {
        "cities": ["Sfax", "Sousse", "Tunis"],
        "country-flag-emoji": "🇹🇳",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Uganda": {
        "cities": ["Kampala", "Kira", "Nansana"],
        "country-flag-emoji": "🇺🇬",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Zambia": {
        "cities": ["Kitwe", "Lusaka", "Ndola"],
        "country-flag-emoji": "🇿🇲",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    "Zimbabwe": {
        "cities": ["Bulawayo", "Chitungwiza", "Harare"],
        "country-flag-emoji": "🇿🇼",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-africa.png",
        "other-terms": [],
        "regions": ["Africa"],
    },
    # REGION: Asia
    "China": {
        "cities": [
            "Beijing",
            "Chengdu",
            "Chongqing",
            "Guangzhou",
            "Hong Kong",
            "Nanjing",
            "Shanghai",
            "Shenzhen",
            "Suzhou",
            "Tianjin",
            "Wuhan",
        ],
        "country-flag-emoji": "🇨🇳",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Chinese"],
        "regions": ["Asia"],
    },
    "India": {
        "cities": [
            "Ahmedabad",
            "Bengaluru",
            "Chennai",
            "Hyderabad",
            "Indore",
            "Jaipur",
            "Kanpur",
            "Kolkata",
            "Kozhikode",
            "Lucknow",
            "Mumbai",
            "Nagpur",
            "New Delhi",
            "Pune",
            "Surat",
        ],
        "country-flag-emoji": "🇮🇳",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": [
            "Andhra Pradesh",
            "Arunachal Pradesh",
            "Assam",
            "Bihar",
            "Chhattisgarh",
            "Goa",
            "Gujarat",
            "Haryana",
            "Himachal Pradesh",
            "Indian",
            "Jharkhand",
            "Karnataka",
            "Kerala",
            "Madhya Pradesh",
            "Maharashtra",
            "Manipur",
            "Meghalaya",
            "Mizoram",
            "Nagaland",
            "Odisha",
            "Punjab",
            "Rajasthan",
            "Sikkim",
            "Tamil Nadu",
            "Telangana",
            "Tripura",
            "Uttar Pradesh",
            "Uttarakhand",
            "West Bengal",
            "india's",
            "india’s",
        ],
        "regions": ["Asia"],
    },
    "Indonesia": {
        "cities": ["Jakarta", "Nusantara"],
        "country-flag-emoji": "🇮🇩",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Indonesian"],
        "regions": ["Asia"],
    },
    "Japan": {
        "cities": ["Chiba", "Fukuoka", "Kawasaki", "Kobe", "Kyoto", "Nagoya", "Osaka", "Saitama", "Tokyo", "Yokohama"],
        "country-flag-emoji": "🇯🇵",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Japanese"],
        "regions": ["Asia"],
    },
    "Malaysia": {
        "cities": ["Kuala Lumpur", "Johor"],
        "country-flag-emoji": "🇲🇾",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Malaysian"],
        "regions": ["Asia"],
    },
    "Pakistan": {
        "cities": ["Islamabad"],
        "country-flag-emoji": "🇵🇰",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": [],
        "regions": ["Asia"],
    },
    "Philippines": {
        "cities": ["Cebu City", "Manila"],
        "country-flag-emoji": "🇵🇭",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Filipino", "Philippine"],
        "regions": ["Asia"],
    },
    "Russia": {
        "cities": [
            "Chelyabinsk",
            "Kazan",
            "Krasnoyarsk",
            "Moscow",
            "Nizhny Novgorod",
            "Novosibirsk",
            "Omsk",
            "Saint Petersburg",
            "Samara",
            "Yekaterinburg",
        ],
        "country-flag-emoji": "🇷🇺",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Rosatom", "Russian"],
        "regions": ["Asia", "Europe"],
    },
    "Singapore": {
        "cities": [],
        "country-flag-emoji": "🇸🇬",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Singaporean"],
        "regions": ["Asia"],
    },
    "South Korea": {
        "cities": ["Busan", "Daegu", "Daejeon", "Incheon", "Seoul"],
        "country-flag-emoji": "🇰🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Korea", "Korean"],
        "regions": ["Asia"],
    },
    "Taiwan": {
        "cities": ["Taipei"],
        "country-flag-emoji": "🇹🇼",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Taiwanese"],
        "regions": ["Asia"],
    },
    "Thailand": {
        "cities": ["Bangkok", "Chiang Mai", "Krung Thep"],
        "country-flag-emoji": "🇹🇭",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Thai"],
        "regions": ["Asia"],
    },
    "Turkey": {
        "cities": ["Ankara", "Antalya", "Bursa", "Istanbul", "Izmir"],
        "country-flag-emoji": "🇹🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Turkish"],
        "regions": ["Asia", "Europe"],
    },
    "Vietnam": {
        "cities": ["Hanoi", "Hà Nội"],
        "country-flag-emoji": "🇻🇳",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Vietnamese"],
        "regions": ["Asia"],
    },
    # REGION: Europe
    "Austria": {
        "cities": ["Graz", "Innsbruck", "Linz", "Salzburg", "Vienna"],
        "country-flag-emoji": "🇦🇹",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Austrian"],
        "regions": ["Europe"],
    },
    "Belgium": {
        "cities": ["Antwerp", "Bruges", "Brussels", "Ghent", "Liege"],
        "country-flag-emoji": "🇧🇪",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": [],
        "regions": ["Europe"],
    },
    "Bulgaria": {
        "cities": [
            "Sofia",
        ],
        "country-flag-emoji": "🇧🇬",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": [
            "Bulgarian",
        ],
        "regions": ["Europe"],
    },
    "Croatia": {
        "cities": [
            "Dubrovnik",
            "Split",
            "Zagreb",
        ],
        "country-flag-emoji": "🇭🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Croatian"],
        "regions": ["Europe"],
    },
    "Cyprus": {
        "cities": ["Nicosia"],
        "country-flag-emoji": "🇨🇾",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Cypriot"],
        "regions": ["Europe"],
    },
    "Czech Republic": {
        "cities": ["Brno", "Liberec", "Ostrava", "Pilsen", "Prague"],
        "country-flag-emoji": "🇨🇿",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Czech", "Czechia"],
        "regions": ["Europe"],
    },
    "Denmark": {
        "cities": ["Aalborg", "Aarhus", "Copenhagen", "København", "Odense"],
        "country-flag-emoji": "🇩🇰",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Danish"],
        "regions": ["Europe", "Nordics"],
    },
    "Estonia": {
        "cities": ["Kohtla-Järve", "Narva", "Pärnu", "Tallinn", "Tartu"],
        "country-flag-emoji": "🇪🇪",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Estonian"],
        "regions": ["Europe"],
    },
    "Finland": {
        "cities": ["Espoo", "Helsinki", "Tampere"],
        "country-flag-emoji": "🇫🇮",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Aalto", "Finnish", "InstituteQ"],
        "regions": ["Europe", "Nordics"],
    },
    "France": {
        "cities": [
            "Bordeaux",
            "Grenoble",
            "Lille",
            "Lyon",
            "Lyons",
            "Marseille",
            "Marseilles",
            "Montpellier",
            "Nantes",
            "Nice",
            "Paris",
            "Strasbourg",
            "Toulouse",
        ],
        "country-flag-emoji": "🇫🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["French"],
        "regions": ["Europe"],
    },
    "Germany": {
        "cities": ["Berlin", "Cologne", "Frankfurt", "Garching", "Hamburg", "Köln", "Munich", "München", "Stuttgart"],
        "country-flag-emoji": "🇩🇪",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Bavaria", "Bavarian", "Fraunhofer", "German ", "German-"],
        "regions": ["Europe"],
    },
    "Greece": {
        "cities": ["Athens"],
        "country-flag-emoji": "🇬🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Greek"],
        "regions": ["Europe"],
    },
    "Hungary": {
        "cities": [
            "Budapest",
        ],
        "country-flag-emoji": "🇭🇺",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": [
            "Hungarian",
            "Magyar",
        ],
        "regions": ["Europe"],
    },
    "Iceland": {
        "cities": ["Akureyri", "Reykjavik", "Reykjavík"],
        "country-flag-emoji": "🇮🇸",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Icelandic"],
        "regions": ["Europe"],
    },
    "Ireland": {
        "cities": ["Cork", "Dublin", "Galway", "Limerick"],
        "country-flag-emoji": "🇮🇪",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Irish"],
        "regions": ["Europe"],
    },
    "Italy": {
        "cities": ["Milan", "Naples", "Palermo", "Rome"],
        "country-flag-emoji": "🇮🇹",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Italian"],
        "regions": ["Europe"],
    },
    "Latvia": {
        "cities": [
            "Riga",
        ],
        "country-flag-emoji": "🇱🇻",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": [
            "Latvian",
        ],
        "regions": ["Europe"],
    },
    "Lithuania": {
        "cities": ["Vilnius"],
        "country-flag-emoji": "🇱🇹",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Lithuanian"],
        "regions": ["Europe"],
    },
    "Luxembourg": {
        "cities": ["Differdange", "Dudelange", "Esch-sur-Alzette", "Luxembourg City", "Pétange"],
        "country-flag-emoji": "🇱🇺",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": [],
        "regions": ["Europe"],
    },
    "Malta": {
        "cities": ["Valletta"],
        "country-flag-emoji": "🇲🇹",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Maltese"],
        "regions": ["Europe"],
    },
    "Netherlands": {
        "cities": ["Amsterdam", "Delft", "Eindhoven", "Rotterdam", "The Hague", "Utrecht"],
        "country-flag-emoji": "🇳🇱",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Dutch"],
        "regions": ["Europe"],
    },
    "Norway": {
        "cities": ["Bergen", "Drammen", "Oslo", "Stavanger", "Trondheim"],
        "country-flag-emoji": "🇳🇴",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Norwegian"],
        "regions": ["Europe", "Nordics"],
    },
    "Poland": {
        "cities": ["Cracow", "Kraków", "Lodz", "Poznan", "Poznań", "Warsaw", "Warszawa", "Wroclaw", "Wrocław", "Łódź"],
        "country-flag-emoji": "🇵🇱",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Polish"],
        "regions": ["Europe"],
    },
    "Portugal": {
        "cities": ["Cascais", "Lisbon", "Porto", "Sintra", "Vila Nova de Gaia"],
        "country-flag-emoji": "🇵🇹",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Portuguese"],
        "regions": ["Europe"],
    },
    "Romania": {
        "cities": ["Bucharest", "Cluj-Napoca", "Constanta", "Iasi", "Timisoara"],
        "country-flag-emoji": "🇷🇴",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Romanian"],
        "regions": ["Europe"],
    },
    "Serbia": {
        "cities": ["Belgrade", "Beograd", "Novi Sad"],
        "country-flag-emoji": "🇷🇸",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Serbian"],
        "regions": ["Europe"],
    },
    "Slovakia": {
        "cities": [
            "Bratislava",
            "Košice",
            "Prešov",
        ],
        "country-flag-emoji": "🇸🇰",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": [
            "Slovak",
        ],
        "regions": ["Europe"],
    },
    "Slovenia": {
        "cities": ["Ljubljana", "Maribor"],
        "country-flag-emoji": "🇸🇮",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Slovenian", "Slovene"],
        "regions": ["Europe"],
    },
    "Spain": {
        "cities": ["Barcelona", "Madrid", "Malaga", "Málaga", "Sevilla", "Valencia", "Zaragoza"],
        "country-flag-emoji": "🇪🇸",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["España"],
        "regions": ["Europe"],
    },
    "Sweden": {
        "cities": ["Gothenburg", "Göteborg", "Malmö", "Stockholm", "Uppsala", "Västerås"],
        "country-flag-emoji": "🇸🇪",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Swedish"],
        "regions": ["Europe", "Nordics"],
    },
    "Switzerland": {
        "cities": ["Basel", "Bern", "Berne", "Geneva", "Genève", "Lausanne", "Zurich", "Zürich"],
        "country-flag-emoji": "🇨🇭",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Swiss"],
        "regions": ["Europe"],
    },
    "UK": {
        "cities": [
            "Belfast",
            "Bournemouth",
            "Brighton",
            "Bristol",
            "Cardiff",
            "Coventry",
            "Edinburgh",
            "Glasgow",
            "Leeds",
            "Leicester",
            "Liverpool",
            "London",
            "Manchester",
            "Newcastle",
            "Newcastle upon Tyne",
            "Nottingham",
            "Sheffield",
            "Southampton",
        ],
        "country-flag-emoji": "🇬🇧",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["NQCC", "UKRI", "United Kingdom"],
        "regions": ["Europe"],
    },
    "Ukraine": {
        "cities": ["Kyiv", "Lviv", "Odesa", "Kiev", "Odessa"],
        "country-flag-emoji": "🇺🇦",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-europe.png",
        "other-terms": ["Ukrainian"],
        "regions": ["Europe"],
    },
    # REGION: Middle East
    "Bahrain": {
        "cities": ["Manama"],
        "country-flag-emoji": "🇧🇭",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Bahraini"],
        "regions": ["Middle East"],
    },
    "Iran": {
        "cities": ["Tehran", "Teheran"],
        "country-flag-emoji": "🇮🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Iranian", "Persia", "Persian"],
        "regions": ["Middle East"],
    },
    "Israel": {
        "cities": ["Haifa", "Jerusalem", "Tel Aviv"],
        "country-flag-emoji": "🇮🇱",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Israeli"],
        "regions": ["Middle East"],
    },
    "Jordan": {
        "cities": [
            "Amman",
            "Aqaba",
        ],
        "country-flag-emoji": "🇯🇴",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": [
            "Jordanian",
        ],
        "regions": ["Middle East"],
    },
    "Kuwait": {
        "cities": ["Al Ahmadi", "Al Farwaniyah", "Hawalli", "Kuwait City", "Salmiya"],
        "country-flag-emoji": "🇰🇼",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Kuwaiti"],
        "regions": ["Middle East"],
    },
    "Lebanon": {
        "cities": [
            "Beirut",
            "Sidon",
            "Tyre",
        ],
        "country-flag-emoji": "🇱🇧",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": [
            "Lebanese",
            "Levant",
        ],
        "regions": ["Middle East"],
    },
    "Oman": {
        "cities": ["Masqat", "Muscat"],
        "country-flag-emoji": "🇴🇲",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Omani"],
        "regions": ["Middle East"],
    },
    "Qatar": {
        "cities": ["Al Khor", "Al Wakrah", "Ar Rayyan", "Doha", "Lusail"],
        "country-flag-emoji": "🇶🇦",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": [],
        "regions": ["Middle East"],
    },
    "Saudi Arabia": {
        "cities": ["Dammam", "Jeddah", "Mecca", "Medina", "Riyadh"],
        "country-flag-emoji": "🇸🇦",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Aramco", "KAUST", "Saudi", "Saudi Aramco"],
        "regions": ["Middle East"],
    },
    "United Arab Emirates": {
        "cities": [],
        "country-flag-emoji": "🇦🇪",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-asia.png",
        "other-terms": ["Abu Dhabi", "Dubai", "UAE"],
        "regions": ["Middle East"],
    },
    # REGION: North America
    "Canada": {
        "cities": [
            "Calgary",
            "Edmonton",
            "Mississauga",
            "Montreal",
            "Montréal",
            "Ottawa",
            "Quebec City",
            "Québec City",
            "Sherbrooke",
            "Toronto",
            "Vancouver",
            "Winnipeg",
        ],
        "country-flag-emoji": "🇨🇦",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": [
            "Alberta",
            "British Columbia",
            "Canadian",
            "Labrador",
            "Manitoba",
            "Newfoundland",
            "Nova Scotia",
            "Ontario",
            "Prince Edward Island",
            "Quebec",
            "Québec",
            "Saskatchewan",
        ],
        "regions": ["North America"],
    },
    "Mexico": {
        "cities": ["Guadalajara", "Mexico City", "Monterrey", "Tijuana"],
        "country-flag-emoji": "🇲🇽",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": ["Mexican"],
        "regions": ["North America"],
    },
    "United States": {
        "cities": [],
        "country-flag-emoji": "🇺🇸",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": [],
        "regions": ["North America"],
    },
    # REGION: Oceania
    "Australia": {
        "cities": ["Adelaide", "Brisbane", "Canberra", "Melbourne", "Perth", "Sydney"],
        "country-flag-emoji": "🇦🇺",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2026/02/flaticon-earth-symbol-with-asia-and-oceania.png",
        "other-terms": ["New South Wales", "Queensland", "Tasmania", "UNSW"],
        "regions": ["Oceania"],
    },
    "New Zealand": {
        "cities": [
            "Auckland",
            "Christchurch",
            "Wellington",
        ],
        "country-flag-emoji": "🇳🇿",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2026/02/flaticon-earth-symbol-with-asia-and-oceania.png",
        "other-terms": [
            "Maori",
        ],
        "regions": ["Oceania"],
    },
    # REGION: South America
    "Argentina": {
        "cities": ["Buenos Aires", "Córdoba"],
        "country-flag-emoji": "🇦🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": ["Argentinian"],
        "regions": ["South America"],
    },
    "Bolivia": {
        "cities": ["La Paz", "Sucre"],
        "country-flag-emoji": "🇧🇴",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": ["Bolivian"],
        "regions": ["South America"],
    },
    "Brazil": {
        "cities": ["Brasilia", "Brasília", "Rio de Janeiro", "Sao Paulo", "São Paulo"],
        "country-flag-emoji": "🇧🇷",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": ["Brazilian"],
        "regions": ["South America"],
    },
    "Chile": {
        "cities": ["Santiago"],
        "country-flag-emoji": "🇨🇱",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": ["Chilean"],
        "regions": ["South America"],
    },
    "Colombia": {
        "cities": ["Bogota", "Bogotá"],
        "country-flag-emoji": "🇨🇴",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": ["Colombian"],
        "regions": ["South America"],
    },
    "Peru": {
        "cities": ["Arequipa", "Lima"],
        "country-flag-emoji": "🇵🇪",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": ["Peruvian"],
        "regions": ["South America"],
    },
    "Uruguay": {
        "cities": ["Montevideo", "Salto"],
        "country-flag-emoji": "🇺🇾",
        "country-icon-link": "https://sutorgroupintelligenceandadvisory.com/wp-content/uploads/2025/10/flaticon-earth-americas.png",
        "other-terms": ["Uruguayan"],
        "regions": ["South America"],
    },
}

eu_countries = [
    "Austria",
    "Belgium",
    "Bulgaria",
    "Croatia",
    "Cyprus",
    "Czech Republic",
    "Denmark",
    "Estonia",
    "Finland",
    "France",
    "Germany",
    "Greece",
    "Hungary",
    "Ireland",
    "Italy",
    "Latvia",
    "Lithuania",
    "Luxembourg",
    "Malta",
    "Netherlands",
    "Poland",
    "Portugal",
    "Romania",
    "Slovakia",
    "Slovenia",
    "Spain",
    "Sweden",
]

middle_east_countries = [
    "Bahrain",
    "Cyprus",
    "Egypt",
    "Iran",
    "Iraq",
    "Israel",
    "Jordan",
    "Kuwait",
    "Lebanon",
    "Oman",
    "Palestine",
    "Qatar",
    "Saudi Arabia",
    "Syria",
    "Turkey",
    "United Arab Emirates",
    "Yemen",
]

african_countries = [
    "Algeria",
    "Angola",
    "Benin",
    "Botswana",
    "Burkina Faso",
    "Burundi",
    "Cabo Verde",
    "Cameroon",
    "Central African Republic",
    "Chad",
    "Comoros",
    "Democratic Republic of the Congo",
    "Republic of the Congo",
    "Djibouti",
    "Egypt",
    "Equatorial Guinea",
    "Eritrea",
    "Eswatini",
    "Ethiopia",
    "Gabon",
    "Gambia",
    "Ghana",
    "Guinea",
    "Guinea-Bissau",
    "Ivory Coast",
    "Kenya",
    "Lesotho",
    "Liberia",
    "Libya",
    "Madagascar",
    "Malawi",
    "Mali",
    "Mauritania",
    "Mauritius",
    "Morocco",
    "Mozambique",
    "Namibia",
    "Niger",
    "Nigeria",
    "Rwanda",
    "Sao Tome and Principe",
    "Senegal",
    "Seychelles",
    "Sierra Leone",
    "Somalia",
    "South Africa",
    "South Sudan",
    "Sudan",
    "Tanzania",
    "Togo",
    "Tunisia",
    "Uganda",
    "Zambia",
    "Zimbabwe",
]

# EMEA (Europe, Middle East, and Africa)
emea_countries = [
    "Albania",
    "Algeria",
    "Andorra",
    "Angola",
    "Armenia",
    "Austria",
    "Azerbaijan",
    "Bahrain",
    "Belarus",
    "Belgium",
    "Benin",
    "Bosnia and Herzegovina",
    "Botswana",
    "Bulgaria",
    "Burkina Faso",
    "Burundi",
    "Cameroon",
    "Cape Verde",
    "Central African Republic",
    "Chad",
    "Comoros",
    "Croatia",
    "Cyprus",
    "Czech Republic",
    "Denmark",
    "Djibouti",
    "Egypt",
    "Equatorial Guinea",
    "Eritrea",
    "Estonia",
    "Eswatini",
    "Ethiopia",
    "Faroe Islands",
    "Finland",
    "France",
    "Gabon",
    "Gambia",
    "Georgia",
    "Germany",
    "Ghana",
    "Gibraltar",
    "Greece",
    "Guernsey",
    "Guinea",
    "Guinea-Bissau",
    "Hungary",
    "Iceland",
    "Iran",
    "Iraq",
    "Ireland",
    "Isle of Man",
    "Israel",
    "Italy",
    "Ivory Coast",
    "Jersey",
    "Jordan",
    "Kenya",
    "Kuwait",
    "Latvia",
    "Lebanon",
    "Lesotho",
    "Liberia",
    "Libya",
    "Liechtenstein",
    "Lithuania",
    "Luxembourg",
    "Macedonia",
    "Madagascar",
    "Malawi",
    "Mali",
    "Malta",
    "Mauritania",
    "Mauritius",
    "Moldova",
    "Monaco",
    "Montenegro",
    "Morocco",
    "Mozambique",
    "Namibia",
    "Netherlands",
    "Niger",
    "Nigeria",
    "Norway",
    "Oman",
    "Palestine",
    "Poland",
    "Portugal",
    "Qatar",
    "Romania",
    "Russia",
    "Rwanda",
    "San Marino",
    "Sao Tome and Principe",
    "Saudi Arabia",
    "Senegal",
    "Serbia",
    "Seychelles",
    "Sierra Leone",
    "Slovakia",
    "Slovenia",
    "Somalia",
    "South Africa",
    "Spain",
    "Sudan",
    "Sweden",
    "Switzerland",
    "Syria",
    "Tanzania",
    "Togo",
    "Tunisia",
    "Turkey",
    "Uganda",
    "Ukraine",
    "United Arab Emirates",
    "United Kingdom",
    "Vatican City",
    "Western Sahara",
    "Yemen",
    "Zambia",
    "Zimbabwe",
]

# NA (North America)
na_countries = [
    "Antigua and Barbuda",
    "Bahamas",
    "Barbados",
    "Belize",
    "Canada",
    "Costa Rica",
    "Cuba",
    "Dominica",
    "Dominican Republic",
    "El Salvador",
    "Grenada",
    "Guatemala",
    "Haiti",
    "Honduras",
    "Jamaica",
    "Mexico",
    "Nicaragua",
    "Panama",
    "Saint Kitts and Nevis",
    "Saint Lucia",
    "Saint Vincent and the Grenadines",
    "Trinidad and Tobago",
    "United States",
]

# LATAM (Latin America)
latam_countries = [
    "Argentina",
    "Belize",
    "Bolivia",
    "Brazil",
    "Chile",
    "Colombia",
    "Costa Rica",
    "Cuba",
    "Dominica",
    "Dominican Republic",
    "Ecuador",
    "El Salvador",
    "French Guiana",
    "Guatemala",
    "Guyana",
    "Haiti",
    "Honduras",
    "Jamaica",
    "Mexico",
    "Nicaragua",
    "Panama",
    "Paraguay",
    "Peru",
    "Saint Kitts and Nevis",
    "Saint Lucia",
    "Saint Vincent and the Grenadines",
    "Suriname",
    "Trinidad and Tobago",
    "Uruguay",
    "Venezuela",
]

# APAC (Asia-Pacific)
apac_countries = [
    "Afghanistan",
    "Australia",
    "Bangladesh",
    "Bhutan",
    "Brunei",
    "Cambodia",
    "China",
    "Cook Islands",
    "Fiji",
    "India",
    "Indonesia",
    "Japan",
    "Kiribati",
    "Laos",
    "Malaysia",
    "Maldives",
    "Marshall Islands",
    "Micronesia",
    "Mongolia",
    "Myanmar",
    "Nepal",
    "New Caledonia",
    "New Zealand",
    "Niue",
    "North Korea",
    "Pakistan",
    "Palau",
    "Papua New Guinea",
    "Philippines",
    "Singapore",
    "Solomon Islands",
    "South Korea",
    "Sri Lanka",
    "Thailand",
    "Timor-Leste",
    "Tonga",
    "Tuvalu",
    "Vanuatu",
    "Vietnam",
]

# APME (Asia-Pacific Middle East)

apme_countries = [
    "Australia",
    "Bahrain",
    "Bangladesh",
    "Bhutan",
    "Brunei",
    "Cambodia",
    "China",
    "Fiji",
    "India",
    "Indonesia",
    "Iran",
    "Iraq",
    "Israel",
    "Japan",
    "Jordan",
    "Kiribati",
    "Kuwait",
    "Laos",
    "Lebanon",
    "Malaysia",
    "Maldives",
    "Marshall Islands",
    "Micronesia",
    "Mongolia",
    "Myanmar (Burma)",
    "Nauru",
    "Nepal",
    "New Zealand",
    "North Korea",
    "Oman",
    "Pakistan",
    "Palau",
    "Papua New Guinea",
    "Philippines",
    "Qatar",
    "Samoa",
    "Saudi Arabia",
    "Singapore",
    "Solomon Islands",
    "South Korea",
    "Sri Lanka",
    "Syria",
    "Taiwan",
    "Thailand",
    "Timor-Leste",
    "Tonga",
    "Tuvalu",
    "United Arab Emirates",
    "Vanuatu",
    "Vietnam",
    "Yemen",
]

# AMER (North, Central, and South America)

amer_countries = [
    "Anguilla",
    "Antigua and Barbuda",
    "Argentina",
    "Aruba",
    "Barbados",
    "Belize",
    "Bermuda",
    "Bolivia",
    "Brazil",
    "British Virgin Islands",
    "Canada",
    "Cayman Islands",
    "Chile",
    "Colombia",
    "Costa Rica",
    "Cuba",
    "Dominica",
    "Dominican Republic",
    "Ecuador",
    "El Salvador",
    "Greenland",
    "Grenada",
    "Guadeloupe",
    "Guatemala",
    "Guyana",
    "Haiti",
    "Honduras",
    "Jamaica",
    "Martinique",
    "Mexico",
    "Montserrat",
    "Netherlands Antilles",
    "Nicaragua",
    "Panama",
    "Paraguay",
    "Peru",
    "Puerto Rico",
    "Saint Kitts and Nevis",
    "Saint Lucia",
    "Saint Pierre and Miquelon",
    "Saint Vincent and the Grenadines",
    "Suriname",
    "The Bahamas",
    "Trinidad and Tobago",
    "Turks and Caicos Islands",
    "U.S. Virgin Islands",
    "United States",
    "Uruguay",
    "Venezuela",
]

oceania_countries = [
    "Australia",
    "New Zealand",
    "Papua New Guinea",
    "Fiji",
    "Solomon Islands",
    "Vanuatu",
    "Samoa",
    "Kiribati",
    "Tonga",
    "Micronesia",
    "Marshall Islands",
    "Palau",
    "Tuvalu",
    "Nauru",
]

nordic_countries = ["Denmark", "Finland", "Iceland", "Norway", "Sweden"]

regions_data = {
    # this is the order they are shown on the chart
    "APAC": {"countries": apac_countries, "name": "Asia-Pacific"},
    "EMEA": {"countries": emea_countries, "name": "Europe, Middle East, and Africa"},
    "LATAM": {"countries": latam_countries, "name": "Latin America"},
    "NA": {"countries": na_countries, "name": "North America"},
    "AMER": {"countries": amer_countries, "name": "North, Central, and South America"},
    "APME": {"countries": apme_countries, "name": "Asia Pacific and Middle East"},
    "EU": {"countries": eu_countries, "name": "European Union"},
    "OC": {"countries": oceania_countries, "name": "Oceania"},
    "Nordics": {"countries": nordic_countries, "name": "Nordics"},
}

regions_counts = dict()  # type: ignore

for _region_abbreviation, _region_data in regions_data.items():
    regions_counts[_region_abbreviation] = dict()
    regions_counts[_region_abbreviation]["count"] = 0
    regions_counts[_region_abbreviation]["countries"] = dict()
    for _country in _region_data["countries"]:
        regions_counts[_region_abbreviation]["countries"][_country] = 0

# U.S. states and territories

us_states_and_territories: dict[str, dict[str, str | list[str]]] = {
    "AL": {"name": "Alabama", "emoji": "🐘", "cities": ["Birmingham", "Huntsville", "Montgomery"], "other-terms": []},
    "AK": {"name": "Alaska", "emoji": "❄️", "cities": ["Anchorage", "Fairbanks", "Juneau"], "other-terms": []},
    "AS": {"name": "American Samoa", "emoji": "🇦🇸", "cities": ["Pago Pago", "Tafuna"], "other-terms": []},
    "AZ": {"name": "Arizona", "emoji": "🌵", "cities": ["Phoenix", "Tucson"], "other-terms": []},
    "AR": {"name": "Arkansas", "emoji": "💎", "cities": ["Fayetteville", "Little Rock"], "other-terms": []},
    "CA": {
        "name": "California",
        "emoji": "🐻",
        "cities": ["Los Angeles", "Sacramento", "San Diego", "San Francisco"],
        "other-terms": [],
    },
    "CO": {"name": "Colorado", "emoji": "🏔️", "cities": ["Boulder", "Colorado Springs", "Denver"], "other-terms": []},
    "CT": {"name": "Connecticut", "emoji": "⚓", "cities": ["Bridgeport", "Hartford", "Stamford"], "other-terms": []},
    "DE": {"name": "Delaware", "emoji": "🐔", "cities": [], "other-terms": []},
    "DC": {"name": "District of Columbia", "emoji": "🏛️", "cities": [], "other-terms": []},
    "FL": {"name": "Florida", "emoji": "☀️", "cities": ["Jacksonville", "Miami", "Orlando", "Tallahassee"], "other-terms": []},
    "GA": {"name": "Georgia", "emoji": "🍑", "cities": ["Atlanta", "Augusta"], "other-terms": []},
    "GU": {"name": "Guam", "emoji": "�", "cities": ["Dededo", "Hagåtña", "Yigo"], "other-terms": []},
    "HI": {"name": "Hawaii", "emoji": "�", "cities": ["Honolulu"], "other-terms": []},
    "ID": {"name": "Idaho", "emoji": "🥔", "cities": ["Boise"], "other-terms": []},
    "IL": {"name": "Illinois", "emoji": "🏙️", "cities": ["Chicago"], "other-terms": ["IQMP"]},
    "IN": {"name": "Indiana", "emoji": "�️", "cities": ["Indianapolis"], "other-terms": []},
    "IA": {"name": "Iowa", "emoji": "�", "cities": ["Cedar Rapids", "Des Moines"], "other-terms": []},
    "KS": {"name": "Kansas", "emoji": "�", "cities": ["Topeka", "Wichita"], "other-terms": []},
    "KY": {"name": "Kentucky", "emoji": "🐎", "cities": ["Louisville"], "other-terms": []},
    "LA": {"name": "Louisiana", "emoji": "🎷", "cities": ["Baton Rouge", "New Orleans", "Shreveport"], "other-terms": []},
    "ME": {"name": "Maine", "emoji": "🦞", "cities": [], "other-terms": []},
    "MD": {"name": "Maryland", "emoji": "�", "cities": ["Annapolis", "Baltimore"], "other-terms": []},
    "MA": {"name": "Massachusetts", "emoji": "🎓", "cities": ["Boston", "Worcester"], "other-terms": []},
    "MI": {"name": "Michigan", "emoji": "🚗", "cities": ["Detroit", "Grand Rapids", "Lansing"], "other-terms": []},
    "MN": {"name": "Minnesota", "emoji": "�", "cities": ["Minneapolis", "St. Paul"], "other-terms": []},
    "MS": {"name": "Mississippi", "emoji": "🎸", "cities": ["Gulfport"], "other-terms": []},
    "MO": {"name": "Missouri", "emoji": "🫏", "cities": ["Kansas City", "St. Louis"], "other-terms": []},
    "MT": {"name": "Montana", "emoji": "⛰️", "cities": ["Billings", "Helena", "Missoula"], "other-terms": []},
    "NE": {"name": "Nebraska", "emoji": "🌾", "cities": ["Omaha"], "other-terms": []},
    "NV": {"name": "Nevada", "emoji": "�", "cities": ["Carson City", "Las Vegas"], "other-terms": []},
    "NH": {"name": "New Hampshire", "emoji": "⛰️", "cities": ["Nashua"], "other-terms": []},
    "NJ": {"name": "New Jersey", "emoji": "🎡", "cities": ["Jersey City", "Newark", "Princeton", "Trenton"], "other-terms": []},
    "NM": {"name": "New Mexico", "emoji": "�️", "cities": ["Albuquerque", "Las Cruces", "Santa Fe"], "other-terms": []},
    "NY": {
        "name": "New York",
        "emoji": "🗽",
        "cities": ["Albany", "Buffalo", "NYC", "New York City", "Rochester", "Stonybrook", "Syracuse"],
        "other-terms": ["Schumer", "Gillibrand", "Hochul"],
    },
    "NC": {"name": "North Carolina", "emoji": "✈️", "cities": ["Raleigh"], "other-terms": []},
    "ND": {"name": "North Dakota", "emoji": "🦬", "cities": ["Bismarck", "Fargo", "Grand Forks"], "other-terms": []},
    "MP": {"name": "Northern Mariana Islands", "emoji": "🇲🇵", "cities": ["Rota", "Saipan", "Tinian"], "other-terms": []},
    "OH": {"name": "Ohio", "emoji": "🌰", "cities": ["Cincinnati", "Cleveland", "Dayton"], "other-terms": []},
    "OK": {"name": "Oklahoma", "emoji": "�️", "cities": ["Oklahoma City", "Tulsa"], "other-terms": []},
    "OR": {"name": "Oregon", "emoji": "�", "cities": ["Portland"], "other-terms": []},
    "PA": {"name": "Pennsylvania", "emoji": "🔔", "cities": ["Harrisburg", "Philadelphia", "Pittsburgh"], "other-terms": []},
    "PR": {"name": "Puerto Rico", "emoji": "�🇷", "cities": ["Bayamón", "San Juan"], "other-terms": []},
    "RI": {"name": "Rhode Island", "emoji": "⚓", "cities": ["Cranston", "Providence"], "other-terms": []},
    "SC": {"name": "South Carolina", "emoji": "⛳", "cities": [], "other-terms": []},
    "SD": {"name": "South Dakota", "emoji": "🗿", "cities": ["Rapid City", "Sioux Falls"], "other-terms": []},
    "TN": {"name": "Tennessee", "emoji": "🎸", "cities": ["Knoxville", "Memphis", "Nashville"], "other-terms": []},
    "TX": {
        "name": "Texas",
        "emoji": "🛢️",
        "cities": ["Austin", "Dallas", "Forth Worth", "Houston", "San Antonio"],
        "other-terms": [],
    },
    "UT": {"name": "Utah", "emoji": "🧗", "cities": ["Salt Lake City"], "other-terms": []},
    "VT": {"name": "Vermont", "emoji": "�", "cities": ["Burlington", "Montpelier"], "other-terms": []},
    "VA": {"name": "Virginia", "emoji": "🏛️", "cities": ["Richmond", "Virginia Beach"], "other-terms": []},
    "VI": {
        "name": "Virgin Islands",
        "emoji": "🇻🇮",
        "cities": ["Charlotte Amalie", "Christiansted", "Cruz Bay"],
        "other-terms": [],
    },
    "WA": {"name": "Washington", "emoji": "🍎", "cities": ["Olympia", "Seattle", "Spokane"], "other-terms": []},
    "WV": {"name": "West Virginia", "emoji": "⛏️", "cities": ["Huntington", "Morgantown"], "other-terms": []},
    "WI": {"name": "Wisconsin", "emoji": "🧀", "cities": ["Green Bay", "Madison", "Milwaukee"], "other-terms": []},
    "WY": {"name": "Wyoming", "emoji": "🤠", "cities": ["Casper", "Cheyenne"], "other-terms": []},
}


us_governors = {
    "Alabama": {"Governor": "Kay Ivey", "Party": "Republican"},
    "Alaska": {"Governor": "Mike Dunleavy", "Party": "Republican"},
    "Arizona": {"Governor": "Katie Hobbs", "Party": "Democratic"},
    "Arkansas": {"Governor": "Sarah Huckabee Sanders", "Party": "Republican"},
    "California": {"Governor": "Gavin Newsom", "Party": "Democratic"},
    "Colorado": {"Governor": "Jared Polis", "Party": "Democratic"},
    "Connecticut": {"Governor": "Ned Lamont", "Party": "Democratic"},
    "Delaware": {"Governor": "Matt Meyer", "Party": "Democratic"},
    "Florida": {"Governor": "Ron DeSantis", "Party": "Republican"},
    "Georgia": {"Governor": "Brian Kemp", "Party": "Republican"},
    "Hawaii": {"Governor": "Josh Green", "Party": "Democratic"},
    "Idaho": {"Governor": "Brad Little", "Party": "Republican"},
    "Illinois": {"Governor": "J.B. Pritzker", "Party": "Democratic"},
    "Indiana": {"Governor": "Eric Holcomb", "Party": "Republican"},
    "Iowa": {"Governor": "Kim Reynolds", "Party": "Republican"},
    "Kansas": {"Governor": "Laura Kelly", "Party": "Democratic"},
    "Kentucky": {"Governor": "Andy Beshear", "Party": "Democratic"},
    "Louisiana": {"Governor": "Jeff Landry", "Party": "Republican"},
    "Maine": {"Governor": "Janet Mills", "Party": "Democratic"},
    "Maryland": {"Governor": "Wes Moore", "Party": "Democratic"},
    "Massachusetts": {"Governor": "Maura Healey", "Party": "Democratic"},
    "Michigan": {"Governor": "Gretchen Whitmer", "Party": "Democratic"},
    "Minnesota": {"Governor": "Tim Walz", "Party": "Democratic"},
    "Mississippi": {"Governor": "Tate Reeves", "Party": "Republican"},
    "Missouri": {"Governor": "Mike Kehoe", "Party": "Republican"},
    "Montana": {"Governor": "Greg Gianforte", "Party": "Republican"},
    "Nebraska": {"Governor": "Jim Pillen", "Party": "Republican"},
    "Nevada": {"Governor": "Joe Lombardo", "Party": "Republican"},
    "New Hampshire": {"Governor": "Kelly Ayotte", "Party": "Republican"},
    "New Jersey": {"Governor": "Phil Murphy", "Party": "Democratic"},
    "New Mexico": {"Governor": "Michelle Lujan Grisham", "Party": "Democratic"},
    "New York": {"Governor": "Kathy Hochul", "Party": "Democratic"},
    "North Carolina": {"Governor": "Josh Stein", "Party": "Democratic"},
    "North Dakota": {"Governor": "Kelly Armstrong", "Party": "Republican"},
    "Ohio": {"Governor": "Mike DeWine", "Party": "Republican"},
    "Oklahoma": {"Governor": "Kevin Stitt", "Party": "Republican"},
    "Oregon": {"Governor": "Tina Kotek", "Party": "Democratic"},
    "Pennsylvania": {"Governor": "Josh Shapiro", "Party": "Democratic"},
    "Rhode Island": {"Governor": "Daniel McKee", "Party": "Democratic"},
    "South Carolina": {"Governor": "Henry McMaster", "Party": "Republican"},
    "South Dakota": {"Governor": "Larry Rhoden", "Party": "Republican"},
    "Tennessee": {"Governor": "Bill Lee", "Party": "Republican"},
    "Texas": {"Governor": "Greg Abbott", "Party": "Republican"},
    "Utah": {"Governor": "Spencer Cox", "Party": "Republican"},
    "Vermont": {"Governor": "Phil Scott", "Party": "Republican"},
    "Virginia": {"Governor": "Glenn Youngkin", "Party": "Republican"},
    "Washington": {"Governor": "Bob Ferguson", "Party": "Democratic"},
    "West Virginia": {"Governor": "Patrick Morrisey", "Party": "Republican"},
    "Wisconsin": {"Governor": "Tony Evers", "Party": "Democratic"},
    "Wyoming": {"Governor": "Mark Gordon", "Party": "Republican"},
}

# Canadian provinces and territories

canadian_provinces_and_territories = {
    "AB": "Alberta",
    "BC": "British Columbia",
    "MB": "Manitoba",
    "NB": "New Brunswick",
    "NL": "Newfoundland and Labrador",
    "NS": "Nova Scotia",
    "NT": "Northwest Territories",
    "NU": "Nunavut",
    "ON": "Ontario",
    "PE": "Prince Edward Island",
    "QC": "Québec",
    "SK": "Saskatchewan",
    "YT": "Yukon",
}


# Australian states and territories

australian_states_and_territories = {
    "ACT": "Australian Capital Territory",
    "NSW": "New South Wales",
    "NT": "Northern Territory",
    "QLD": "Queensland",
    "SA": "South Australia",
    "TAS": "Tasmania",
    "VIC": "Victoria",
    "WA": "Western Australia",
}

# German federal states

german_states = {
    "BW": "Baden-Württemberg",
    "BY": "Bavaria",
    "BE": "Berlin",
    "BB": "Brandenburg",
    "HB": "Bremen",
    "HH": "Hamburg",
    "HE": "Hesse",
    "NI": "Lower Saxony",
    "MV": "Mecklenburg-Vorpommern",
    "NW": "North Rhine-Westphalia",
    "RP": "Rhineland-Palatinate",
    "SL": "Saarland",
    "SN": "Saxony",
    "ST": "Saxony-Anhalt",
    "SH": "Schleswig-Holstein",
    "TH": "Thuringia",
}

# French modern region names

french_current_regions_names = [
    "Auvergne-Rhône-Alpes",
    "Bourgogne-Franche-Comté",
    "Bretagne",
    "Centre-Val de Loire",
    "Corse",
    "Grand Est",
    "Hauts-de-France",
    "Île-de-France",
    "Normandie",
    "Nouvelle-Aquitaine",
    "Occitanie",
    "Pays de la Loire",
    "Provence-Alpes-Côte d'Azur",
]

# Indian states and union territories

indian_states_and_union_territories = [
    "Andaman and Nicobar Islands",
    "Andhra Pradesh",
    "Arunachal Pradesh",
    "Assam",
    "Bihar",
    "Chandigarh",
    "Chhattisgarh",
    "Dadra & Nagar Haveli and Daman & Diu",
    "Delhi",
    "Goa",
    "Gujarat",
    "Haryana",
    "Himachal Pradesh",
    "Jammu and Kashmir",
    "Jharkhand",
    "Karnataka",
    "Kerala",
    "Ladakh",
    "Lakshadweep",
    "Madhya Pradesh",
    "Maharashtra",
    "Manipur",
    "Meghalaya",
    "Mizoram",
    "Nagaland",
    "Odisha",
    "Puducherry",
    "Punjab",
    "Rajasthan",
    "Sikkim",
    "Tamil Nadu",
    "Telangana",
    "Tripura",
    "Uttar Pradesh",
    "Uttarakhand",
    "West Bengal",
]

# Japanese prefectures

japanese_prefectures = [
    "Aichi",
    "Akita",
    "Aomori",
    "Chiba",
    "Ehime",
    "Fukui",
    "Fukuoka",
    "Fukushima",
    "Gifu",
    "Gunma",
    "Hiroshima",
    "Hokkaido",
    "Hyogo",
    "Ibaraki",
    "Ishikawa",
    "Iwate",
    "Kagawa",
    "Kagoshima",
    "Kanagawa",
    "Kochi",
    "Kumamoto",
    "Kyoto",
    "Mie",
    "Miyagi",
    "Miyazaki",
    "Nagano",
    "Nagasaki",
    "Nara",
    "Niigata",
    "Oita",
    "Okayama",
    "Okinawa",
    "Osaka",
    "Saga",
    "Saitama",
    "Shiga",
    "Shimane",
    "Shizuoka",
    "Tochigi",
    "Tokushima",
    "Tokyo",
    "Tottori",
    "Toyama",
    "Wakayama",
    "Yamagata",
    "Yamaguchi",
    "Yamanashi",
]

tokyo_special_wards = [
    "Chiyoda-ku",
    "Chuo-ku",
    "Minato-ku",
    "Shinjuku-ku",
    "Bunkyo-ku",
    "Taito-ku",
    "Sumida-ku",
    "Koto-ku",
    "Shinagawa-ku",
    "Meguro-ku",
    "Ota-ku",
    "Setagaya-ku",
    "Shibuya-ku",
    "Nakano-ku",
    "Suginami-ku",
    "Toshima-ku",
    "Kita-ku",
    "Arakawa-ku",
    "Itabashi-ku",
    "Nerima-ku",
    "Adachi-ku",
    "Katsushika-ku",
    "Edogawa-ku",
]

# cspell:enable


def get_region_abbreviations():
    return [region for region in regions_data]


def increment_country_in_regions(country):
    if country == "USA":
        country = "United States"
    elif country == "UK":
        country = "United Kingdom"
    elif country == "UAE":
        country = "United Arab Emirates"
    elif country == "The Netherlands":
        country = "Netherlands"
        print("Replace 'The Netherlands' with 'Netherlands'")

    found_a_region = False

    for region, countries_in_region in regions_counts.items():
        if country in countries_in_region["countries"]:
            found_a_region = True
            regions_counts[region]["count"] += 1
            countries_in_region["countries"][country] += 1

    if not found_a_region:
        print(f"Country {country} is not in any region.")


def print_countries_in_regions_counts():
    for region_name, data in regions_counts.items():
        title = f"{region_name}: {data['count']}"
        print(f"\n{title}\n{len(title) * '-'}")
        for country, country_count in data["countries"].items():
            if country_count > 0:
                print(f"  {country}: {country_count}")


city_counts_by_country: dict[str, dict[str, int]] = dict()


def increment_city_count(address: str):
    country = get_country(address)
    city = get_city(address)

    if country == "Japan" and city in tokyo_special_wards:
        city = f"{city} - Tokyo"

    if country not in city_counts_by_country:
        city_counts_by_country[country] = dict()

    if city not in city_counts_by_country[country]:
        city_counts_by_country[country][city] = 1
    else:
        city_counts_by_country[country][city] += 1


def get_city_count(country: str):
    cities: list[str] = []
    counts: list[int] = []

    if country in city_counts_by_country:
        city_counts_by_country[country] = dict(sorted(city_counts_by_country[country].items()))

        for city, count in city_counts_by_country[country].items():
            cities.append(city)
            counts.append(count)

    return cities, counts


# -----------------------------------------------------------------------------
# Canada
# -----------------------------------------------------------------------------


Canada_province_counts: dict[str, int] = dict()
Canada_city_counts: dict[str, int] = dict()

for _province in canadian_provinces_and_territories.values():
    Canada_province_counts[_province] = 0


def increment_Canada_province_count(address):
    province = get_state_or_province(address)
    assert province in canadian_provinces_and_territories.values()
    Canada_province_counts[province] += 1  # type: ignore


def get_Canada_province_counts():
    provinces = []
    counts = []
    for _province, _count in Canada_province_counts.items():
        if _count != 0:
            provinces.append(_province)
            counts.append(_count)
    return provinces, counts


def increment_Canada_city_count(address):
    city = get_city(address)
    if city not in Canada_city_counts:
        Canada_city_counts[city] = 1
    else:
        Canada_city_counts[city] += 1


# -----------------------------------------------------------------------------
# France
# -----------------------------------------------------------------------------

France_region_counts: dict[str, int] = dict()
France_city_counts: dict[str, int] = dict()

for _region in french_current_regions_names:
    France_region_counts[_region] = 0


def increment_France_region_count(address):
    region = get_state_or_province(address)
    if region not in french_current_regions_names:
        raise ValueError(f"Unknown French region: {region}")
    France_region_counts[region] += 1  # type: ignore


def get_France_region_counts():
    regions = []
    counts = []
    for _region, _count in France_region_counts.items():
        if _count != 0:
            regions.append(_region)
            counts.append(_count)
    return regions, counts


def increment_France_city_count(address):
    city = get_city(address)
    if city not in France_city_counts:
        France_city_counts[city] = 1
    else:
        France_city_counts[city] += 1


# -----------------------------------------------------------------------------
# Germany
# -----------------------------------------------------------------------------

Germany_state_counts: dict[str, int] = dict()
Germany_city_counts: dict[str, int] = dict()

for _state in german_states.values():
    Germany_state_counts[_state] = 0


def increment_Germany_state_count(address):
    state = get_state_or_province(address)
    if state not in german_states.values():
        raise ValueError(f"Unknown German state: {state}")
    Germany_state_counts[state] += 1  # type: ignore


def get_Germany_state_counts():
    states = []
    counts = []
    for _state, _count in Germany_state_counts.items():
        if _count != 0:
            states.append(_state)
            counts.append(_count)
    return states, counts


def increment_Germany_city_count(address):
    city = get_city(address)
    if city not in Germany_city_counts:
        Germany_city_counts[city] = 1
    else:
        Germany_city_counts[city] += 1


# -----------------------------------------------------------------------------
# Italy
# -----------------------------------------------------------------------------

Italy_city_counts: dict[str, int] = dict()


def increment_Italy_city_count(address):
    city = get_city(address)
    if city not in Italy_city_counts:
        Italy_city_counts[city] = 1
    else:
        Italy_city_counts[city] += 1


# -----------------------------------------------------------------------------
# Japan
# -----------------------------------------------------------------------------

Japan_city_counts: dict[str, int] = dict()


def increment_Japan_city_count(address):
    city = get_city(address)

    if city in tokyo_special_wards:
        city = f"{city} - Tokyo"

    if city not in Japan_city_counts:
        Japan_city_counts[city] = 1
    else:
        Japan_city_counts[city] += 1


# -----------------------------------------------------------------------------
# Spain
# -----------------------------------------------------------------------------

Spain_city_counts: dict[str, int] = dict()


def increment_Spain_city_count(address):
    city = get_city(address)
    if city not in Spain_city_counts:
        Spain_city_counts[city] = 1
    else:
        Spain_city_counts[city] += 1


# -----------------------------------------------------------------------------
# UK
# -----------------------------------------------------------------------------

UK_city_counts: dict[str, int] = dict()


def increment_UK_city_count(address):
    city = get_city(address)
    if city not in UK_city_counts:
        UK_city_counts[city] = 1
    else:
        UK_city_counts[city] += 1


# -----------------------------------------------------------------------------
# US
# -----------------------------------------------------------------------------

us_state_counts = dict()

for state_values in us_states_and_territories.values():
    us_state_counts[state_values["name"]] = 0


def increment_US_state_count(address):
    state = get_state_or_province(address)
    # assert state in us_states_and_territories.values()
    us_state_counts[state] += 1  # type: ignore


def get_US_state_counts():
    states = []
    counts = []
    for _state, _count in us_state_counts.items():
        if _count != 0:
            states.append(_state)
            counts.append(_count)
    return states, counts


# -----------------------------------------------------------------------------
# General city or region/state increments
# -----------------------------------------------------------------------------


def increment_state_province_or_local_region(address):
    increment_city_count(address)
    country = get_country(address)

    if country == "USA":
        increment_US_state_count(address)

    elif country == "Canada":
        increment_Canada_province_count(address)

    elif country == "France":
        increment_France_region_count(address)

    elif country == "Germany":
        increment_Germany_state_count(address)


def get_max_companies_in_a_state_or_province():
    m = max(us_state_counts.values())
    m = max(m, *Canada_province_counts.values())
    m = max(m, *France_region_counts.values())
    m = max(m, *Germany_state_counts.values())
    return m


def get_regions_and_counts():
    """Return all region abbreviations and their company counts, sorted case-insensitively.

    Returns:
        tuple[list[str], list[int]]: A tuple of (regions, counts) where regions is a
        sorted list of region abbreviation strings and counts is the corresponding list
        of company counts for each region.
    """
    regions = sorted(list(regions_counts.keys()), key=str.casefold)
    counts = [regions_counts[region]["count"] for region in regions]

    return (regions, counts)


def write_region_html_appendix(id_, appendix_title, yattag_tag, yattag_text):
    with yattag_tag("h2", id=id_):
        yattag_text(appendix_title)

    for region_abbreviation, region_data in regions_data.items():
        with yattag_tag("p", klass="region-name"):
            yattag_text(f"{region_data['name']} ({region_abbreviation})")

        with yattag_tag("p", klass="country-list"):
            yattag_text(", ".join(region_data["countries"]))


def write_region_word_appendix(id_, appendix_title, word_document, text=None):
    # pylint: disable=W0212
    heading = word_document.add_heading(appendix_title, 1)
    run = heading.runs[0]
    docx_tools.add_bookmark_for_id(id_, run)

    if text is not None:
        docx_tools.convert_html_to_word(text, word_document)

    docx_tools.insert_horizontal_rule(word_document)

    # Create a custom region style
    region_style = word_document.styles.add_style("RegionAppendix", 1)
    region_style.font.name = "Aptos"
    region_style.font.bold = True
    region_style.font.italic = False
    region_style.font.size = word_document.styles["Normal"].font.size

    sorted_by_region = dict(sorted(regions_data.items(), key=lambda item: item[0]))

    for region_abbreviation, region_data in sorted_by_region.items():
        # p = word_document.add_paragraph(
        #     f"{region_data['name']} ({region_abbreviation})", style="RegionAppendix"
        # )

        p = word_document.add_paragraph(f"{region_abbreviation} – {region_data['name']}", style="RegionAppendix")
        docx_tools.mark_index_entry(region_abbreviation, p)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True

        p = word_document.add_paragraph(", ".join(region_data["countries"]))
        p.paragraph_format.left_indent = Pt(18)
    # pylint: enable=W0212


def get_city(address_):
    while "  " in address_:
        address_ = address_.replace("  ", " ")

    if ", " not in address_:
        return address_

    address_parts = address_.split(", ")
    return address_parts[0].strip()


def get_country(address_):
    while "  " in address_:
        address_ = address_.replace("  ", " ")

    if ", " not in address_:
        return address_

    address_parts = address_.split(", ")
    return address_parts[-1].strip()


def get_state_or_province(address_):
    # for simplicity, we call them states
    # could also be provinces and French regions, for example

    while "  " in address_:
        address_ = address_.replace("  ", " ")

    if ", " not in address_:
        return address_

    address_parts = address_.split(", ")

    # Handle cities that are their own states or provinces

    if address_parts[0] in ["Berlin"]:
        return address_parts[0]

    if len(address_parts) > 2:
        state = address_parts[-2].strip()
        country = address_parts[-1].strip()

        if country == "USA" and state in us_states_and_territories:
            return us_states_and_territories[state]["name"]

        if country == "Canada" and state in canadian_provinces_and_territories:
            return canadian_provinces_and_territories[state]

        if country == "Australia" and state in australian_states_and_territories:
            return australian_states_and_territories[state]

        return state

    return None
