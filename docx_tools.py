# general purpose tools

# cspell:ignore autofit oxml powerpnt rels taskkill twips winword

from copy import copy

import docx
from bs4 import BeautifulSoup
from common_data import HARVARD_CRIMSON, SANS_SERIF_FONT
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# -------------------------------------------------------------------------------------------------
# Hyperlinks
# -------------------------------------------------------------------------------------------------


def add_bookmark_for_id(id_, run):
    # Create bookmark elements

    # pylint: disable=W0212
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), id_)  # Unique ID (use a different one for each bookmark)
    bookmark_start.set(qn("w:name"), id_)  # Bookmark name

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), id_)  # Same ID as bookmarkStart

    # Insert bookmark before and after the text
    run._r.insert(0, bookmark_start)  # Before text
    run._r.addnext(bookmark_end)  # After text
    # pylint: enable=W0212


def add_hyperlink(paragraph, text_, url, is_italic=False):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # pylint: disable=W0212

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
    hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(docx.oxml.shared.OxmlElement("w:r"), paragraph)
    new_run.text = text_

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    new_run.font.underline = False

    if is_italic:
        new_run.italic = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)

    # pylint: disable=W0212

    return new_run


def add_bookmark(paragraph, bookmark_name, bookmark_id):
    # pylint: disable=W0212
    # Create bookmarkStart element
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), bookmark_name)
    # Create bookmarkEnd element
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))
    # Insert bookmarkStart at the beginning, bookmarkEnd at the end
    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)
    # pylint: disable=W0212


def add_hyperlink_to_bookmark(paragraph, link_text, bookmark_name, tooltip=None, is_italic=False):
    """Creates a hyperlink to an internal bookmark."""

    get_or_create_hyperlink_style(paragraph.part.document)

    # pylint: disable=W0212
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    if tooltip is not None:
        hyperlink.set(qn("w:tooltip"), tooltip)

    run = paragraph.add_run()
    rPr = run._element.get_or_add_rPr()
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")  # use the hyperlink style in word
    rPr.append(rStyle)

    new_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = link_text
    new_run.append(text)
    hyperlink.append(new_run)
    run._element.append(hyperlink)
    # pylint: enable=W0212

    if is_italic:
        run.italic = True
    return run


def get_or_create_hyperlink_style(d_):
    """
    If this document had no hyperlinks so far, the builtin Hyperlink style will likely be missing and we need
    to add it. There's no predefined value, different Word versions define it differently. This version is how
    Word 2019 defines it in the default theme, excluding a theme reference.
    """

    if "Hyperlink" not in d_.styles:
        if "Default Character Font" not in d_.styles:
            ds = d_.styles.add_style("Default Character Font", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True)
            ds.element.set(docx.oxml.shared.qn("w:default"), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d_.styles.add_style("Hyperlink", docx.enum.style.WD_STYLE_TYPE.CHARACTER, True)
        hs.base_style = d_.styles["Default Character Font"]
        hs.unhide_when_used = True

        # hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        # ROYAL_BLUE = "#0044B9"
        hs.font.color.rgb = docx.shared.RGBColor(0x00, 0x44, 0xB9)

        hs.font.underline = True
        del hs

    remove_hyperlink_underline(d_)

    return "Hyperlink"


def remove_hyperlink_underline(doc):
    """Removes the underline from the Hyperlink style in a Word document."""
    styles = doc.styles
    hyperlink_style = styles.element.xpath('w:style[@w:styleId="Hyperlink"]')[0]

    # Check if rPr exists, if not create it.
    rPr = hyperlink_style.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        hyperlink_style.append(rPr)

    # Check if u element exists, remove it if it does.
    u = rPr.find(qn("w:u"))
    if u is not None:
        rPr.remove(u)


# -------------------------------------------------------------------------------------------------
# Table of Contents
# -------------------------------------------------------------------------------------------------


def insert_word_toc(document):
    # Code for making Table of Contents
    # Source: https://stackoverflow.com/questions/18595864/python-create-a-table-of-contents-with-python-docx-lxml

    # style = document.styles["TOC1"]

    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    # pylint: disable=W0212

    fldChar = OxmlElement("w:fldChar")  # creates a new element
    fldChar.set(qn("w:fldCharType"), "begin")  # sets attribute on element

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click to update field."

    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    # p_element = paragraph._p
    run.font.underline = False
    # pylint: enable=W0212


# -------------------------------------------------------------------------------------------------
# Footers
# -------------------------------------------------------------------------------------------------

# Function to add a page number field in the footer


def add_page_number(paragraph):
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()

    # pylint: disable=W0212
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    # pylint: enable=W0212


def set_triple_footer(section, left_text, center_text):
    section.footer.is_linked_to_previous = False
    footer = section.footer

    content_width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips

    # Create a table in the footer with three columns
    table = footer.add_table(rows=1, cols=3, width=content_width)
    table.autofit = True

    # Access cells and add text
    cell_left = table.cell(0, 0)
    cell_left.width = Inches(1.0)
    cell_center = table.cell(0, 1)
    cell_center.width = Inches(5.0)
    cell_right = table.cell(0, 2)
    cell_right.width = Inches(1.0)

    # Add paragraphs to cells and align text
    paragraph_left = cell_left.paragraphs[0]
    paragraph_left.add_run(left_text)
    paragraph_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

    paragraph_center = cell_center.paragraphs[0]
    paragraph_center.add_run(center_text)
    paragraph_center.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # page number goes on the right
    paragraph_right = cell_right.paragraphs[0]
    add_page_number(paragraph_right)
    paragraph_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # set the foot font size
    FOOTER_PTS = 8

    for para in (paragraph_left, paragraph_center, paragraph_right):
        paragraph_format = para.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # Set exact line spacing
        paragraph_format.line_spacing = Pt(FOOTER_PTS)

        for run in para.runs:
            run.font.size = Pt(FOOTER_PTS)


# -------------------------------------------------------------------------------------------------
# Styles
# -------------------------------------------------------------------------------------------------


def copy_style(source_document, target_document, style_name: str):
    # https://stackoverflow.com/questions/78733174/how-can-i-use-styles-from-an-existing-docx-file-in-my-new-document
    source_style_names = [style.name for style in source_document.styles]
    target_style_names = [style.name for style in target_document.styles]
    if style_name in source_style_names:
        source_style = source_document.styles[style_name]
        if style_name in target_style_names:
            target_document.styles[style_name].delete()
        target_document.styles.element.append(copy(source_style.element))


# -------------------------------------------------------------------------------------------------
# Tables
# -------------------------------------------------------------------------------------------------


def set_word_table_heading_cell_colors(cell):
    # pylint: disable=W0212
    tcPr = cell._element.get_or_add_tcPr()
    # pylint: enable=W0212

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), HARVARD_CRIMSON)  # Red color
    tcPr.append(shd)

    for p in cell.paragraphs:
        # p.paragraph_format.space_before = Pt(3)
        # p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)  # white
            run.font.name = SANS_SERIF_FONT


def set_word_table_cell_colors(cell, row_index):
    if row_index % 2 == 0:  # even numbered row
        # pylint: disable=W0212
        tcPr = cell._element.get_or_add_tcPr()
        # pylint: enable=W0212

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "#EEEEEE")
        tcPr.append(shd)


# -------------------------------------------------------------------------------------------------
# Page breaks
# -------------------------------------------------------------------------------------------------


def insert_page_break(word_document):
    assert word_document is not None

    word_document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# -------------------------------------------------------------------------------------------------
# Section breaks
# -------------------------------------------------------------------------------------------------


def insert_section_break(word_document):
    assert word_document is not None

    return word_document.add_section(WD_SECTION.NEW_PAGE)


def insert_two_column_section(word_document):
    assert word_document is not None

    section = word_document.add_section(WD_SECTION.CONTINUOUS)

    # pylint: disable=W0212

    sectPr = section._sectPr  # Access the section's XML

    # Find or create the <w:cols> element
    cols = sectPr.xpath("./w:cols")
    if not cols:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    else:
        cols = cols[0]

    # Set the number of columns and spacing
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "500")  # Space between columns (in twentieths of a point)
    # pylint: enable=W0212


def insert_one_column_section(word_document):
    assert word_document is not None

    section = word_document.add_section(WD_SECTION.CONTINUOUS)

    # pylint: disable=W0212

    sectPr = section._sectPr  # Access the section's XML

    # Find or create the <w:cols> element
    cols = sectPr.xpath("./w:cols")
    if not cols:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    else:
        cols = cols[0]

    # Set the number of columns and spacing
    cols.set(qn("w:num"), "1")
    # pylint: enable=W0212


# -------------------------------------------------------------------------------------------------
# Conversion of simple HTML to Word
# -------------------------------------------------------------------------------------------------


def convert_html_to_word(html_string, word_document):
    # if there is no HTML markup, just add the text as a paragraph

    # while "  " in html_string:
    #     html_string = html_string.replace("  ", " ")

    html_string = html_string.replace("<p> ", "<p>")
    # html_string = html_string.replace(" <p> ", "<p>")
    html_string = html_string.replace(" </p>", "</p>")
    # html_string = html_string.replace("</p> ", "</p>")
    html_string = html_string.replace("<p> </p>", "<p></p>")

    new_paragraph = False
    is_italic = False
    is_bold = False
    in_ul = False
    in_ol = False

    def walk_html(element, p_):
        nonlocal new_paragraph, is_italic, is_bold, in_ul, in_ol

        for child in element.contents:
            if child.name:
                # print(f"'{child.name}'")

                if child.name == "dt":
                    new_paragraph = True
                    p_ = word_document.add_paragraph()
                    p_.paragraph_format.space_after = Pt(0)
                    p_.paragraph_format.keep_with_next = True
                    is_bold = True
                    walk_html(child, p_)
                    is_bold = False

                elif child.name == "dd":
                    new_paragraph = True
                    p_ = word_document.add_paragraph()
                    p_.paragraph_format.left_indent = Pt(18)
                    walk_html(child, p_)

                elif child.name == "br":
                    run = p_.runs[-1]
                    run.add_break(WD_BREAK.LINE)

                elif child.name == "p":
                    p_ = word_document.add_paragraph()
                    new_paragraph = True

                    classes = child.get("class")

                    if isinstance(classes, list):
                        if "align-center" in classes:
                            p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif "align-left" in classes:
                            p_.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif "align-right" in classes:
                            p_.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif "align-justify" in classes:
                            p_.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        # else:
                        #     print(f"Paragraph classes: {classes}")

                        if "hanging-indent" in classes:
                            p_.paragraph_format.left_indent = Pt(36)
                            p_.paragraph_format.first_line_indent = Pt(-18)

                    walk_html(child, p_)

                elif child.name == "h2":
                    new_paragraph = True
                    p_ = word_document.add_heading(child.string, level=1)

                elif child.name == "h3":
                    new_paragraph = True
                    p_ = word_document.add_heading(child.string, level=2)

                elif child.name == "h4":
                    new_paragraph = True
                    p_ = word_document.add_heading(child.string, level=3)

                elif child.name == "a":
                    new_paragraph = False
                    the_link = child.get("href")
                    if isinstance(the_link, str):
                        if p_ is None:
                            p_ = word_document.add_paragraph()
                            new_paragraph = True

                        if not the_link.startswith("#"):
                            add_hyperlink(p_, child.string, the_link, is_italic=is_italic)
                        else:
                            add_hyperlink_to_bookmark(p_, child.string, the_link[1:], is_italic=is_italic)
                    else:
                        run = p_.add_run(child.string)
                        if is_italic:
                            run.italic = True

                elif child.name == "em" or child.name == "i":
                    new_paragraph = False
                    is_italic = True
                    walk_html(child, p_)
                    is_italic = False

                elif child.name == "strong" or child.name == "b":
                    new_paragraph = False
                    is_bold = True
                    walk_html(child, p_)
                    is_bold = False

                elif child.name == "ul":
                    new_paragraph = False
                    in_ul = True
                    walk_html(child, p_)
                    in_ul = False

                elif child.name == "ol":
                    new_paragraph = False
                    in_ol = True
                    walk_html(child, p_)
                    in_ol = False

                elif child.name == "li":
                    new_paragraph = True
                    if in_ol:
                        p_ = word_document.add_paragraph("", style="List Number")
                    else:
                        p_ = word_document.add_paragraph("", style="List Bullet")
                    walk_html(child, p_)

                else:
                    walk_html(child, p_)

            elif isinstance(child, str) and child.strip():
                the_text = child.string
                if p_ is None:
                    p_ = word_document.add_paragraph()
                    new_paragraph = True

                if new_paragraph:
                    the_text = the_text.lstrip()

                new_paragraph = False

                run = p_.add_run(the_text)

                if is_italic:
                    run.italic = True

                if is_bold:
                    run.bold = True

    if "<" not in html_string:
        p = word_document.add_paragraph()
        p.add_run(html_string)
    else:
        walk_html(BeautifulSoup(html_string, "html.parser"), None)


# -------------------------------------------------------------------------------------------------
# Index
# -------------------------------------------------------------------------------------------------


def mark_index_entry(entry, paragraph, subentry=None):
    if subentry is not None:
        subentry_code = f":{subentry}"
    else:
        subentry_code = ""

    run = paragraph.add_run()
    r = run._r

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    # instrText.text = ' XE "%s" ' % (entry)  # type: ignore

    if entry[0].isalnum():
        instrText.text = f' XE "{entry}{subentry_code}" '  # type: ignore
    else:
        instrText.text = f' XE "{entry};{entry[1:]}{subentry_code}" '  # type: ignore

    r.append(instrText)

    run = paragraph.add_run()
    r = run._r

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "separate")
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r

    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)


def insert_word_index(document):
    # Code for making an Index
    # Source: https://stackoverflow.com/questions/18595864/python-create-a-table-of-contents-with-python-docx-lxml

    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    # pylint: disable=W0212

    fldChar = OxmlElement("w:fldChar")  # creates a new element
    fldChar.set(qn("w:fldCharType"), "begin")  # sets attribute on element

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")  # sets attribute on element
    instrText.text = 'INDEX \\c 2 \\h "A" \\z "1033" }'  # type: ignore

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click to update field."  # type: ignore

    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

    # p_element = paragraph._p
    run.font.underline = False
    # pylint: enable=W0212


def insert_index(document):
    insert_section_break(document)

    heading = document.add_heading("Index", 1)
    run = heading.runs[0]
    add_bookmark_for_id("section-index", run)

    insert_word_index(document)
