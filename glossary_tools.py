"""
Formatting glossary with yattag
"""

# cspell:ignore addnext asis klass oxml OxmlElement qn rangle tagtext
# cspell:ignore vert yattag rdquo langle lsquo rsquo

import json
from datetime import date

import docx_tools
import os_tools
import yattag
from common_data import SANS_SERIF_FONT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# import sys


GLOSSARY_FILE = "../../data/quantum-glossary.json"

with open(GLOSSARY_FILE, "rt", encoding="utf8") as json_file:
    glossary = json.load(json_file)


entity_substitutions = {
    "&langle": "\u27e8",
    "&ldquo;": "\u201c",
    "&lsquo;": "\u2018",
    "&rangle;": "\u232a",
    "&rdquo;": "\u201d",
    "&rsquo;": "\u2019",
    "&vert;": "|",
}

glossary_intro_text = [
    "This glossary contains common terms and definitions from the areas of quantum computing quantum networking, quantum sensing, and quantum communications. It is a work in progress.",
    "Most of the definitions come from Perplexity and the sources are noted.",
]


def create_html_glossary_file(file_name):
    source_links = {"Perplexity": "https://www.perplexity.ai/"}

    doc, tag, text = yattag.Doc().tagtext()

    with tag("p"):
        today = date.today().strftime("%A, %B %d, %Y").replace(" 0", " ")
        with tag("em"):
            text(f"Last updated {today}.")

    for line in glossary_intro_text:
        with tag("p"):
            doc.asis(line)

    doc.stag("hr")

    for term, content in glossary.items():
        with tag("p"):
            term_ = term
            for entity, unicode in entity_substitutions.items():
                term_ = term_.replace(entity, unicode)

            with tag("strong"):
                text(term_)
        with tag("p", style="margin-left: 18pt;"):
            definition_text = " ".join(content["definition"]).strip()
            while "  " in definition_text:
                definition_text = definition_text.replace("  ", " ")
            for entity, unicode in entity_substitutions.items():
                definition_text = definition_text.replace(entity, unicode)
            doc.asis(definition_text)
            with tag("em"):
                text(" Source: ")
            if content["source"] in source_links:
                with tag("a", href=source_links[content["source"]], target="_blank", rel="noopener"):
                    text(content["source"])
            else:
                text(content["source"])

        if "dwq-reference" in content:
            with tag("p", style="margin-left: 18pt;"):
                text(f"Discussed in section {content['dwq-reference']} of ")
                with tag("a", href="https://amzn.to/4eylDSZ", target="_blank", rel="noopener"):
                    with tag("em"):
                        text("Dancing with Qubits, Second Edition")
                text(".")

    result = yattag.indent(doc.getvalue())
    # result = doc.getvalue()

    result = result.replace("</strong>\n      <a", "</strong> <a")

    with open(file_name, "wt", encoding="utf-8") as output_html_file:
        print(result, file=output_html_file)


def html_format_glossary_term(term, yattag_doc, yattag_tag, yattag_text):
    try:
        term_data = glossary[term]
    except KeyError as exc:
        os_tools.terminating_error(f"Glossary term '{term}' is not in glossary: {exc}")

    term_id = get_glossary_id(term)

    with yattag_tag("p", klass="glossary-term", id=term_id):
        yattag_doc.asis(term)

    with yattag_tag("p", klass="glossary-definition"):
        yattag_doc.asis(" ".join(term_data["definition"]))
        yattag_text(" (")
        with yattag_tag("em"):
            yattag_text("Source: ")

        if term_data["source"] == "Perplexity":
            with yattag_tag("a", href="https://www.perplexity.ai/", target="_blank"):
                yattag_text(term_data["source"])
        else:
            yattag_text(term_data["source"])
        yattag_text(")")


def word_format_glossary_terms(terms, word_document):
    # Create a custom region style
    glossary_style = word_document.styles.add_style("GlossaryAppendix", 1)
    glossary_style.font.name = SANS_SERIF_FONT
    glossary_style.font.bold = True
    glossary_style.font.italic = False
    glossary_style.font.size = word_document.styles["Normal"].font.size

    for term in terms:
        try:
            term_data = glossary[term]
        except KeyError as exc:
            os_tools.terminating_error(f"Glossary term '{term}' is not in glossary: {exc}")

        term_id = get_glossary_id(term)

        term_text = term
        for entity, new_expression in entity_substitutions.items():
            term_text = term_text.replace(entity, new_expression)
        if "&" in term_text:
            print(f"Warning: '&' in glossary term '{term_text}'")

        p = word_document.add_paragraph(term_text, style="GlossaryAppendix")

        if term_text.endswith(" modality"):
            docx_tools.mark_index_entry("modality:" + term_text.replace(" modality", ""), p)
        elif term_text.endswith(" paradigm"):
            docx_tools.mark_index_entry("paradigm:" + term_text.replace(" paradigm", ""), p)

        docx_tools.mark_index_entry(term_text, p)

        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True

        run = p.runs[0]

        # Create bookmark elements
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), term_id)  # Unique ID (use a different one for each bookmark)
        bookmark_start.set(qn("w:name"), term_id)  # Bookmark name

        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), term_id)  # Same ID as bookmarkStart

        # Insert bookmark before and after the text
        # pylint: disable=W0212
        run._r.insert(0, bookmark_start)  # Before text
        run._r.addnext(bookmark_end)  # After text
        # pylint: enable=W0212

        glossary_text = " ".join(term_data["definition"])

        for entity, new_expression in entity_substitutions.items():
            glossary_text = glossary_text.replace(entity, new_expression)
        if "&" in glossary_text:
            print(f"Warning: '&' in glossary text '{glossary_text}'")

        docx_tools.convert_html_to_word(glossary_text, word_document)
        p = word_document.paragraphs[-1]

        # p = word_document.add_paragraph(" ".join(term_data["definition"]))
        p.paragraph_format.left_indent = Pt(18)

        p.add_run(" (")
        run = p.add_run("Source: ")
        run.italic = True
        if term_data["source"] == "Perplexity":
            docx_tools.add_hyperlink(p, term_data["source"], "https://www.perplexity.ai/")
        else:
            p.add_run(term_data["source"])
        p.add_run(")")

        if "notes" in term_data:
            p = word_document.add_paragraph(" ".join(term_data["notes"]))
            p.paragraph_format.left_indent = Pt(18)
            run = p.runs[0]
            run.italic = True


def get_glossary_id(term: str) -> str:
    term = term.replace("&vert;", "|")
    term = term.replace("&rangle;", ">")
    term = term.replace(" ", "-")
    return f"glossary-{term.casefold()}"


def write_glossary_word_appendix(id_, appendix_title, word_document):
    heading = word_document.add_heading(appendix_title, 1)
    run = heading.runs[0]
    docx_tools.add_bookmark_for_id(id_, run)

    p = word_document.add_paragraph()
    p.add_run("This appendix provides a reference for the technical terms used throughout this document.")

    docx_tools.insert_horizontal_rule(word_document)

    # # Create bookmark elements
    # # pylint: disable=W0212
    # bookmark_start = OxmlElement("w:bookmarkStart")
    # bookmark_start.set(qn("w:id"), id_)  # Unique ID (use a different one for each bookmark)
    # bookmark_start.set(qn("w:name"), id_)  # Bookmark name

    # bookmark_end = OxmlElement("w:bookmarkEnd")
    # bookmark_end.set(qn("w:id"), id_)  # Same ID as bookmarkStart

    # # Insert bookmark before and after the text
    # run._r.insert(0, bookmark_start)  # Before text
    # run._r.addnext(bookmark_end)  # After text
    # # pylint: enable=W0212


"""
Give me a 5-sentence English description of the quantum computing term "".
Do not include any links or references. Do not mention use cases or applications.
Do not include matrices in mathematical form. Do not include links or references.
"""

if __name__ == "__main__":
    create_html_glossary_file("test-html-glossary.html")
